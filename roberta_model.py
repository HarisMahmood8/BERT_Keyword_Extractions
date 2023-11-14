from docx import Document
from transformers import AutoModelForSequenceClassification, AutoTokenizer
import torch
import openpyxl
import nltk
from nltk.corpus import stopwords
from nltk import word_tokenize
from test_keyword_read import read_keyword_topics, export_to_excel


def analyze_sentiment(text, model, tokenizer):
    inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True)
    outputs = model(**inputs)
    logits = outputs.logits
    sentiment_score = logits[0][1] - logits[0][0]
    return sentiment_score.item()


def get_categories(keywords_list, categories):
    category_list = set()
    for keyword in keywords_list:
        for category in categories:
            if keyword in categories.get(category):
                category_list.add(category)
    return category_list


def run_sentiment(file):
    model_name = "mrm8488/distilroberta-finetuned-financial-news-sentiment-analysis"
    model = AutoModelForSequenceClassification.from_pretrained(model_name)
    tokenizer = AutoTokenizer.from_pretrained(model_name)

    # company = input("Enter name of company: ")
    # number = int(input("Enter the quarter number: "))
    # quarter = f"Q{number}"
    # doc = Document(f"{company}_{quarter}.docx")
    doc = Document(f'{file}.docx')
    # output_file = f"sentiment_results_{company}_{quarter}.xlsx"
    output_file = f'sentiment_results_{file}.xlsx'

    nltk.download('punkt')
    nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))

    workbook = openpyxl.load_workbook("keywords_topics.xlsx")
    categories_dict = read_keyword_topics(workbook["Key Words_Topics_Streamlined"])
    keywords = [item for sublist in categories_dict.values() for item in sublist]

    result_list = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        text_sentences = text.split(". ")

        for sentence in text_sentences:
            sentiment = analyze_sentiment(sentence, model, tokenizer)
            text_words = sentence.split()

            found_keywords = [keyterm for keyterm in keywords if any(keyword in text_words and word_tokenize(
                keyword)[0].isalpha() and keyword.lower() not in stop_words for keyword in keyterm.split())]

            if found_keywords:
                categories_list = get_categories(found_keywords, categories_dict)
                result_list.append([text, float(sentiment), ', '.join(
                    found_keywords), ', '.join(categories_list)])

    export_to_excel(result_list, output_file)

    workbook.close()

    print(f"Sentiment analysis and keyword extraction completed. Results saved in '{output_file}'.")


if __name__ == "__main__":
    file_name = input("Enter call document file name, without the extension: ")
    run_sentiment(file_name)
